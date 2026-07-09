from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

REPORT_FILE = 'Pengujian_Black_Box_Website_Gurau_Tenda_Kost.docx'

sections = [
    {
        'title': '1. Pengujian Navbar',
        'rows': [
            {'scenario': 'Verifikasi semua link menu Beranda, Sewa Tenda, Kost Putri, Testimoni, Hubungi Kami', 'expected': 'Link mengarah ke halaman yang benar sesuai route dan anchor #contact.', 'actual': 'Semua route dan anchor tersedia pada Blade dan route web.php.', 'note': 'Sesuai'},
            {'scenario': 'Hover effect pada menu link', 'expected': 'Link berubah warna/underline saat hover.', 'actual': 'CSS nav-links a:hover menerapkan warna keemasan dan garis bawah.', 'note': 'Sesuai'},
            {'scenario': 'Toggle dark/light mode', 'expected': 'Tema berubah dan tersimpan di localStorage.', 'actual': 'toggleTheme() mengubah class light-mode dan menyimpan localStorage.', 'note': 'Sesuai'},
            {'scenario': 'Hamburger mobile menu buka/tutup', 'expected': 'Klik hamburger menampilkan nav pada layar kecil.', 'actual': 'toggleMenu() menambahkan kelas open; CSS mobile menampilkan nav-links.open.', 'note': 'Sesuai'},
            {'scenario': 'Tombol Panel Admin', 'expected': 'ADMIN membuka login modal, setelah login menunjukkan panel admin.', 'actual': 'ADMIN membuka modal; login sukses menampilkan panel; jika belum login tombol membuka modal.', 'note': 'Sesuai'},
        ]
    },
    {
        'title': '2. Pengujian Beranda',
        'rows': [
            {'scenario': 'Hero section menampilkan judul, teks, dan tombol CTA', 'expected': 'Hero tampil lengkap dengan teks dan tombol Hubungi Kami.', 'actual': 'Section hero ada dengan judul, subtext, dan tombol anchor.', 'note': 'Sesuai'},
            {'scenario': 'Stats counter animasi', 'expected': 'Statistik menghitung ke angka final.', 'actual': 'animasiCounter() dan animasiRating() menghitung angka 4.8★, 4.2★, 20+.', 'note': 'Sesuai'},
            {'scenario': 'Card Keahlian Kami ditampilkan', 'expected': 'Empat item keahlian hadir di hero card.', 'actual': 'Hero expertise card berisi Tenda, Kost Putri, Keamanan, Informasi.', 'note': 'Sesuai'},
            {'scenario': 'Scroll reveal atas & bawah', 'expected': 'Section scroll-reveal muncul dengan animasi slide naik/turun.', 'actual': 'IntersectionObserver menambahkan kelas reveal-visible saat enter viewport.', 'note': 'Sesuai'},
            {'scenario': 'Tombol Hubungi Kami bekerja', 'expected': 'Scroll ke section kontak.', 'actual': 'Tombol mengarah ke #contact di halaman.', 'note': 'Sesuai'},
            {'scenario': 'Foto grid beranda dan lightbox foto', 'expected': 'Klik foto membuka lightbox dengan navigasi.', 'actual': 'Lightbox global menginisialisasi images dari hero dan gallery.', 'note': 'Sesuai'},
        ]
    },
    {
        'title': '3. Pengujian Sewa Tenda',
        'rows': [
            {'scenario': 'Semua jenis tenda tersedia', 'expected': 'Dropdown memuat Biasa, Semi VIP, VIP, Balon, Sentris.', 'actual': 'Semua opsi ada pada select tentType.', 'note': 'Sesuai'},
            {'scenario': 'Validasi minimal unit Balon', 'expected': 'Balon otomatis minimum 2 unit dengan peringatan.', 'actual': 'onTentTypeChange() menaikkan unit ke 2 dan menampilkan balonWarning.', 'note': 'Sesuai'},
            {'scenario': 'Pilihan ukuran tenda mempengaruhi perhitungan', 'expected': 'Ukuran dapat dipilih dan muncul di breakdown.', 'actual': 'Ukuran tersimpan di breakdown tetapi tidak mempengaruhi total harga.', 'note': 'Tidak Sesuai - ukuran hanya label, bukan faktor harga'},
            {'scenario': 'Warna swatch tampil untuk jenis tenda yang ada pilihan', 'expected': 'Warna dekorasi muncul sebagai swatch; warna baru tersedia jika ditambahkan.', 'actual': 'Swatch ada untuk biasa, semivip, vip; balon/sentris tidak menyediakan opsi warna.', 'note': 'Tidak Sesuai - warna kosong untuk Balon/Sentris, opsi baru tidak terdaftar'},
            {'scenario': 'Pilihan kursi dan jumlah kursi muncul saat dibutuhkan', 'expected': 'Kursi none menyembunyikan qty, pilihan lain menampilkan qty.', 'actual': 'chairQtyRow toggle sesuai pilihan kursi.', 'note': 'Sesuai'},
            {'scenario': 'Panggung dapat dipilih dan menambah estimasi harga', 'expected': 'Pemilihan panggung menambah biaya 350 ribu.', 'actual': 'hitungHarga() menambahkan HARGA_PANGGUNG saat ada.', 'note': 'Sesuai'},
            {'scenario': 'Meja dapat dipilih dan jumlah meja diperbolehkan', 'expected': 'Meja dan jumlahnya muncul dan dihitung.', 'actual': 'mejaQtyRow toggle dan harga meja dihitung.', 'note': 'Sesuai'},
            {'scenario': 'Kalkulasi harga real-time', 'expected': 'Total dan breakdown terupdate saat pilihan berubah.', 'actual': 'hitungHarga() memperbarui total/ breakdown pada setiap perubahan.', 'note': 'Sesuai'},
            {'scenario': 'Konfirmasi via WhatsApp menyimpan data pemesanan lalu membuka chat', 'expected': 'Pesanan disimpan dan WhatsApp dibuka.', 'actual': 'fetch POST dilakukan, lalu WhatsApp dibuka di finally tanpa cek error.', 'note': 'Tidak Sesuai - membuka WhatsApp meski penyimpanan gagal'},
            {'scenario': 'FAQ accordion berfungsi buka/tutup', 'expected': 'FAQ dapat dibuka/ditutup dengan ikon berubah.', 'actual': 'initFAQAccordion() mengendalikan state open dan aria-expanded.', 'note': 'Sesuai'},
        ]
    },
    {
        'title': '4. Pengujian Kost Putri',
        'rows': [
            {'scenario': 'Info harga & fasilitas Kost Putri ditampilkan', 'expected': 'Harga, fasilitas, note, dan tombol WhatsApp terlihat.', 'actual': 'Section kost menampilkan harga, fasilitas, dan button WA.', 'note': 'Sesuai'},
            {'scenario': 'Hover effect fasilitas aktif', 'expected': 'Fasilitas mempunyai hover effect.', 'actual': 'CSS .facility-item:hover mengubah border dan bayangan.', 'note': 'Sesuai'},
            {'scenario': 'Lightbox foto Kost buka/navigasi/tutup/scroll', 'expected': 'Klik thumbnail membuka lightbox, navigasi prev/next, escape, dan swipe.', 'actual': 'Lightbox overlay tersedia dengan tombol prev/next/close dan swipe support.', 'note': 'Sesuai'},
            {'scenario': 'Cek ketersediaan kamar mengirim AJAX dan menampilkan status', 'expected': 'Menampilkan tersedia/penuh sesuai bulan.', 'actual': 'cekKetersediaan() mengambil route /cek-kamar dan menampilkan pesan.', 'note': 'Sesuai'},
            {'scenario': 'Tombol WhatsApp Kost Putri bekerja', 'expected': 'Membuka chat WhatsApp info kost.', 'actual': 'Link WA statis ada di tombol.', 'note': 'Sesuai'},
            {'scenario': 'Section Keunggulan Kami hadir dan animasi muncul', 'expected': 'Tiga advantage card tampil dengan efek scroll.', 'actual': 'Advantage cards ditampilkan dan observer menambahkan kelas visible.', 'note': 'Sesuai'},
            {'scenario': 'Fitur Tambah AC Kost Putri mempengaruhi harga', 'expected': 'Pilihan AC menambah harga display +3.000.000.', 'actual': 'selectAC() mengubah angka harga tetapi tidak terhubung ke pemesanan backend.', 'note': 'Tidak Sesuai - hanya visual display, tidak disimpan/pesanan'}
        ]
    },
    {
        'title': '5. Pengujian Testimoni',
        'rows': [
            {'scenario': 'Filter tab Tenda/Kost bekerja', 'expected': 'Review sesuai kategori tampil, kategori lain tersembunyi.', 'actual': 'filterUlasan() mengubah display card berdasarkan data-cat.', 'note': 'Sesuai'},
            {'scenario': 'Animasi rating bar & counter', 'expected': 'Bar persentase dan angka rating animasi saat muncul.', 'actual': 'animasiRating() memulai animasi bar dan angka.', 'note': 'Sesuai'},
            {'scenario': 'Carousel auto-scroll looping', 'expected': 'Carousel bergerak otomatis dan loop tanpa henti.', 'actual': 'autoScrollStep() scroll terus dan loop via clone cards.', 'note': 'Sesuai'},
            {'scenario': 'Pause on hover dan drag manual', 'expected': 'Hover menghentikan auto-scroll; drag menggeser slide.', 'actual': 'wrapper event listeners pause auto-scroll dan pointer drag bekerja.', 'note': 'Sesuai'},
            {'scenario': 'Navigasi panah dan dots bekerja', 'expected': 'Klik panah/dot pindah slide dan update indikator.', 'actual': 'scrollReview() dan scrollToCard() mengubah index dan dot aktif.', 'note': 'Sesuai'},
        ]
    },
    {
        'title': '6. Pengujian Admin Panel',
        'rows': [
            {'scenario': 'Login admin berhasil dengan kredensial valid', 'expected': 'Session admin dibuat, panel muncul.', 'actual': 'login() memvalidasi config/admin dan menyimpan session admin_login; JS mengatur localStorage dan membuka panel.', 'note': 'Sesuai'},
            {'scenario': 'Login admin gagal dengan kredensial salah', 'expected': 'Mengembalikan error dan menolak akses.', 'actual': 'login() mengembalikan status 401 dengan pesan gagal.', 'note': 'Sesuai'},
            {'scenario': 'Toggle ketersediaan kamar mengubah badge dan data', 'expected': 'Switch mengubah status tersedia/penuh.', 'actual': 'onToggleKetersediaan() memperbarui dataKamar dan badge.', 'note': 'Sesuai'},
            {'scenario': 'Ubah jumlah kamar berhasil', 'expected': 'Menambah/kurangi jumlah kamar dan otomatis set tersedia jika >0.', 'actual': 'adminUbahQty() memperbarui jumlah dan status tersedia.', 'note': 'Sesuai'},
            {'scenario': 'Simpan perubahan admin mengirim data', 'expected': 'POST data status kamar dan pesan sukses tampil.', 'actual': 'simpanDataAdmin() mengirim JSON ke /admin/simpan-status; pesan sukses muncul.', 'note': 'Sesuai'},
            {'scenario': 'History pesanan memuat daftar pemesanan', 'expected': 'Data history tampil pada tab History Pesan.', 'actual': 'muatHistoryPemesanan() fetch /admin/pemesanan dan render list.', 'note': 'Sesuai'},
            {'scenario': 'Logout admin berfungsi', 'expected': 'Session admin dihapus dan panel tersembunyi.', 'actual': 'adminLogout() memanggil /admin/logout lalu menghapus localStorage dan sembunyikan panel.', 'note': 'Sesuai'},
            {'scenario': 'Session persistence antar halaman', 'expected': 'Admin tetap tampil antar reload jika masih login.', 'actual': 'JS membaca localStorage adminLoggedIn dan menampilkan panel, tetapi backend session tidak divalidasi sebelum render.', 'note': 'Tidak Sesuai - persistence dilakukan secara client-side, risiko mismatch sesi'},
            {'scenario': 'Bulk action Semua Tersedia/Penuh', 'expected': 'Tombol bulk action tersedia dan bekerja.', 'actual': 'Fitur tidak ditemukan di kode.', 'note': 'Tidak Sesuai'},
            {'scenario': 'Ubah status pesanan di admin history', 'expected': 'Dapat mengganti status pemesanan.', 'actual': 'Fitur tidak ditemukan di kode.', 'note': 'Tidak Sesuai'},
        ]
    },
    {
        'title': '7. Pengujian Cetak Invoice',
        'rows': [
            {'scenario': 'Tombol cetak invoice di halaman invoice', 'expected': 'Klik memanggil window.print()', 'actual': 'Button print memanggil window.print().', 'note': 'Sesuai'},
            {'scenario': 'Download PDF invoice tersedia', 'expected': 'Tombol membuka route download PDF.', 'actual': 'Link mengarah ke route /admin/invoice/{id}/pdf.', 'note': 'Sesuai'},
            {'scenario': 'Invoice memuat jenis tenda/kamar, harga, jumlah, total', 'expected': 'Seluruh field tampil lengkap.', 'actual': 'Invoice menampilkan jenis, jumlah, ukuran, warna, fasilitas, total; harga per item ditampilkan sebagai \'-\'.', 'note': 'Tidak Sesuai - kolom harga item tidak diisi secara detail'},
            {'scenario': 'Format invoice rapi tanpa data kosong/error', 'expected': 'Halaman invoice terformat baik dan tidak menampilkan nilai kosong.', 'actual': 'Invoice rapi namun harga item berisi simbol \'-\' yang mengindikasikan detail harga tidak lengkap.', 'note': 'Tidak Sesuai'},
        ]
    },
    {
        'title': '8. Pengujian Responsivitas Mobile',
        'rows': [
            {'scenario': 'Navbar dan hamburger responsif', 'expected': 'Hamburger tampil dan nav mobile dapat dibuka.', 'actual': 'CSS mobile menampilkan hamburger dan open-nav.' , 'note': 'Sesuai'},
            {'scenario': 'Layout hero menyesuaikan ukuran layar kecil', 'expected': 'Hero content stacked/tertata di mobile.', 'actual': 'Terdapat CSS mobile untuk layout responsive.', 'note': 'Sesuai'},
            {'scenario': 'Foto grid dan lightbox mobile terbaca', 'expected': 'Lightbox bekerja di layar kecil.', 'actual': 'Lightbox touch/swipe support ada untuk mobile.', 'note': 'Sesuai'},
            {'scenario': 'Carousel swipe manual pada mobile', 'expected': 'Pengguna dapat menggeser kartu review dengan drag.', 'actual': 'Event pointer supports manual drag dan swipe.', 'note': 'Sesuai'},
            {'scenario': 'Admin panel responsif mobile', 'expected': 'Panel tampil full width dan tombol cukup kecil.', 'actual': 'Admin panel melebar ke full width pada mobile dengan tombol ukuran lebih kecil.', 'note': 'Sesuai'},
            {'scenario': 'Ukuran tombol minimal 48px untuk akses sentuh', 'expected': 'Tombol interaktif memiliki target sentuh minimal 48px.', 'actual': 'Beberapa tombol qty berukuran 34px sehingga tidak memenuhi 48px.', 'note': 'Tidak Sesuai'},
        ]
    },
    {
        'title': '9. Pengujian Kontak & Lokasi',
        'rows': [
            {'scenario': 'Informasi kontak dan jam operasional ditampilkan', 'expected': 'Tiga kartu kontak tampil lengkap.', 'actual': 'Section kontak menampilkan telepon, jam, dan lokasi.', 'note': 'Sesuai'},
            {'scenario': 'Floating WhatsApp button tersedia', 'expected': 'Tombol WhatsApp mengambang di semua halaman beranda.', 'actual': 'Link WA floating ada di footer section Kontak.', 'note': 'Sesuai'},
            {'scenario': 'Embed Google Maps menampilkan lokasi', 'expected': 'Dua iframe maps ditampilkan.', 'actual': 'Dua iframe untuk tenda dan kost tersedia.', 'note': 'Sesuai'},
            {'scenario': 'Anchor Hubungi Kami scroll ke section kontak', 'expected': 'Tombol hero membawa ke #contact.', 'actual': 'Hero button tujuan #contact.', 'note': 'Sesuai'},
        ]
    }
]


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = 'Arial'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
style.font.size = Pt(11)

# Title
p = doc.add_paragraph()
p.alignment = 1
run = p.add_run('Pengujian Black Box Website Gurau Tenda & Kost')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Arial'

for section_data in sections:
    doc.add_paragraph()
    h = doc.add_paragraph(section_data['title'])
    h.style = doc.styles['Normal']
    h.runs[0].bold = True
    h.runs[0].font.size = Pt(12)
    h.runs[0].font.name = 'Arial'

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['No', 'Skenario Pengujian', 'Hasil yang Diharapkan', 'Hasil Pengujian / Keterangan']
    for idx, cell in enumerate(hdr_cells):
        cell.text = headers[idx]
        paragraph = cell.paragraphs[0]
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        paragraph.runs[0].font.name = 'Arial'
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'B7D9FF')
        cell._tc.get_or_add_tcPr().append(shading)
    for idx, row in enumerate(section_data['rows'], 1):
        cells = table.add_row().cells
        cells[0].text = str(idx)
        cells[1].text = row['scenario']
        cells[2].text = row['expected']
        cells[3].text = f"{row['actual']} ({row['note']})"
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10.5)

# Summary

num_scenarios = sum(len(s['rows']) for s in sections)
num_ok = sum(1 for s in sections for r in s['rows'] if r['note'] == 'Sesuai')
num_not_ok = sum(1 for s in sections for r in s['rows'] if r['note'] != 'Sesuai')

p = doc.add_paragraph()
p.add_run('Ringkasan Pengujian').bold = True
p.paragraph_format.space_before = Pt(12)

table = doc.add_table(rows=4, cols=2)
for i, (label, value) in enumerate([
    ('Total skenario diuji', str(num_scenarios)),
    ('Jumlah Sesuai', str(num_ok)),
    ('Jumlah Tidak Sesuai', str(num_not_ok)),
    ('Catatan', 'Beberapa fitur ada tetapi tidak sepenuhnya terimplementasi atau detail harga invoice tidak lengkap.'),
]):
    cells = table.rows[i].cells
    cells[0].text = label
    cells[1].text = value
    for cell in cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10.5)


# Save report

doc.save(REPORT_FILE)
print(f'Report created: {REPORT_FILE}')
